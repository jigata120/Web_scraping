import json
from docx import Document
from docx.shared import Pt

def add_paragraph(doc, title, text, level, list_type=None, indent_level=0):
    if title:
        title = title.strip()

        if level == 1:
            doc.add_paragraph("\n")
            doc.add_paragraph("\n")
            p = doc.add_heading(title, level=1)
        elif list_type:
            if list_type == 'unordered':
                p = doc.add_paragraph(title, style='ListBullet')
            elif list_type == 'ordered':
                p = doc.add_paragraph(title, style='ListNumber')
        else:
            if level == 2:
                p = doc.add_heading(title, level=2)
            elif level == 3:
                p = doc.add_heading(title, level=3)
            elif level == 4:
                p = doc.add_heading(title, level=4)
            else:
                p = doc.add_paragraph(title)

        p.paragraph_format.left_indent = Pt(12 * indent_level)

    if text:
        p = doc.add_paragraph(text)
        p.style.font.size = Pt(12)
        p.paragraph_format.left_indent = Pt(12 * indent_level)

def process_node(doc, node, indent_level=0):
    title = node.get('title', '').strip()
    text = node.get('text', '')
    level = node.get('level', 0)

    if title.lower() in ['unordered', 'ordered']:
        list_type = title.lower()
        children = node.get('children', [])
        for child in children:
            add_paragraph(doc, child.get('title', ''), '', child.get('level', 0), list_type, indent_level)
    else:
        add_paragraph(doc, title, text, level, None, indent_level)
        for child in node.get('children', []):
            process_node(doc, child, indent_level + 1)

def json_to_docx(json_data, output_file='output.docx'):
    data = json.loads(json_data)
    doc = Document()

    for item in data.get('objects', []):
        process_node(doc, item)

    doc.save(output_file)

def read_json_from_file(file_path):
    with open(file_path, 'r') as file:
        return file.read()

if __name__ == "__main__":
    json_file_path = 'output.json'
    json_data = read_json_from_file(json_file_path)
    json_to_docx(json_data, 'output.docx')
